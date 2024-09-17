import os
import json
from docx import Document
from django.conf import settings

def userrequirement_generate_docx(context, project_id):
    document = Document()

    # Add project information
    project = context['project']
    document.add_heading('Project Details', level=1)
    document.add_paragraph(f"Project Name: {project.name}")
    document.add_paragraph(f"Project Description: {project.description}")
    document.add_paragraph(f"Project Type: {project.project_type}")

    # Add user responses
    document.add_heading('User Responses', level=1)
    for response in context['user_responses']:
        document.add_paragraph(f"Question: {response.question.text}")
        if response.question.question_type == 'combo' and response.selected_option:
            document.add_paragraph(f"Answer: {response.selected_option.text}")
        if response.question.question_type == 'checkbox' and response.selected_options.exists():
            document.add_paragraph("Answers: ")
            for option in response.selected_options.all():
                document.add_paragraph(f"- {option.text}")

    # Add UI requirements
    ui_requirements = context['ui_requirements']
    document.add_heading('UI Requirements', level=1)
    document.add_paragraph(f"Font: {ui_requirements.font}")
    document.add_paragraph(f"Font Color: {ui_requirements.font_color}")
    document.add_heading('Color Palette', level=2)
    for color in ui_requirements.color_palette:
        color_text = f"rgb({color[0]}, {color[1]}, {color[2]})"
        document.add_paragraph(color_text)

    # Add technical requirements
    document.add_heading('Technical Requirements', level=1)
    for requirement in context['technical_requirements']:
        document.add_paragraph(f"Development Framework: {requirement.development_framework}")
        document.add_paragraph(f"CMS: {requirement.cms}")

    # Save the document
    docx_file_path = os.path.join(settings.STATIC_ROOT, f'documents\project_{project_id}_details.docx')
    document.save(docx_file_path)

    return docx_file_path
