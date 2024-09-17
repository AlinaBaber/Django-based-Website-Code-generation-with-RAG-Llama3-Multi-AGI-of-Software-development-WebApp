from huggingface_hub import login
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import textwrap
import os
#####################
import os
import shutil
import torch
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline
import torch
from docx import Document
import os
import json
import os
import re
from softwaredevelopmentapi.model_loader import llm_model, tokenizer
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch


def load_all_models():
    # Clears cache if you're using CUDA

    global llm_model, tokenizer
    # global img_to_text_pipe
    # Load your model here, e.g., Hugging Face model
    device = "cuda" if torch.cuda.is_available() else "cpu"
    # tokenizer = AutoTokenizer.from_pretrained("meta-llama/Llama-2-7b-chat-hf", use_fast=True, device=device)
    model_name = "meta-llama/Meta-Llama-3-8B-Instruct"
    # model_name="meta-llama/Llama-2-7b-chat-hf"
    llm_model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
                                                     use_auth_token="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")
    terminators = [
        tokenizer.eos_token_id,
        tokenizer.convert_tokens_to_ids("<|eot_id|>")
    ]


class SoftwareDeveloperAgent:
    def __init__(self, model_name="meta-llama/Meta-Llama-3-8B-Instruct",
                 tokenizer_name="meta-llama/Meta-Llama-3-8B-Instruct"):
        # self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
        #                                               use_auth_token="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")
        # self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
        #                                                  use_auth_token="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")
        device = "cuda" if torch.cuda.is_available() else "cpu"
        self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
                                                       use_auth_token="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")
        self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
                                                          use_auth_token="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")
        terminators = [
            tokenizer.eos_token_id,
            tokenizer.convert_tokens_to_ids("<|eot_id|>")
        ]

    def read_pdf(self, file_path):
        loader = PyPDFLoader(file_path)
        data = loader.load()

        return data

    def read_docx(self, file_path):
        loader = Docx2txtLoader(file_path)
        data = loader.load()

        return data

    def read_txt(self, file_path):
        loader = TextLoader(file_path)
        data = loader.load()

        return data

    def add_text_to_docx(self, file_path, text, output_path):
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(output_path)

    def generate_response_from_document(self, file_path, query, project_info):
        # if len(project_info) > 0:
        #    self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
        # """Generate a response based on a user query from a document."""
        # Determine the file type and extract text
        if file_path.lower().endswith('.pdf'):
            document_text = self.read_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            document_text = self.read_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            document_text = self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")

        streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

        text_pipeline = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            max_new_tokens=2000,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=streamer,
        )

        # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

        llm = HuggingFacePipeline(pipeline=text_pipeline, model_kwargs={"temperature": 0.7})

        # Step 2: Split the text into manageable chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        texts = text_splitter.split_documents(document_text)

        # Step 3: Generate embeddings for the texts
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={"device": "cuda"}
        )
        db = Chroma.from_documents(texts, embeddings)

        SYSTEM_PROMPT = f"""
       You are a knowledgeable AI Assistant, responsible for providing code to query . Your responses should be thorough, clear, and aligned with standard object oriented Programming Django python code standards. Always provide Django python code syntax only, while ensuring the code is accurate and relevant to provided project in the document.
Always clarify the context if needed. {query}
        """
        prompt = PromptTemplate(template="""You are a highly knowledgeable Chatbot, dedicated to providing accurate and detailed informative answers to user queries. Please don't generate irrelevant or extra information which the user did not ask you to generate.    
Context: {context}
User: {question}
Chatbot:""", input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt},
        )

        # Execute a query (you can customize this part as needed)
        query_result = qa_chain(query)

        # Execute the query
        #         query_result, source_documents = qa_chain.run(query)

        # print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

        # Extract the answer from the 'response' field using string manipulation
        # Assuming the answer always follows "Answer:" and ends at the end of the string
        response_content = query_result['result']
        answer_prefix = "Chatbot:"
        answer_start_index = response_content.find(answer_prefix)
        if answer_start_index != -1:
            answer = response_content[answer_start_index + len(answer_prefix):].strip()
            print(answer)
            return answer
        else:
            print("No answer found in the response.")
            return response_content

    def software_developer_agent_template(self, file_path, output_path, base_dir, projectid, ui_data):

        query = "Generate the Clothing detailed  website HTML code according to provided sample html in document, Ensure that the generated code strictly adheres to the information provided without adding any extra content or features not mentioned in the document. Provide detailed and clear answers, focusing solely on the specified requirements. you should mention html file name. you should nav should have link of other linked pages to that html, should add web based images as example."
        path = file_path
        #        sections_queries = {
        #          "html": {
        #              "query": f"Write static html single page website with (.html files,style.css,script.js) code for the project based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
        #            "path": file_path}}
        project_info = {}
        #        project_details = ""
        #        for section, details in sections_queries.items():
        print("query", query)
        # query = f"For Project {projectname}, {details.get('query')}"
        response = self.generate_response_from_document(path, query, project_info)
        # if section == "Django Directory":
        #    self.create_directory_structure_from_text(response, base_dir)
        # Ensure the section key is formatted correctly
        section_key = "index"
        # project_info[section_key] = response
        # Create a new .docx document
        doc = Document()
        doc.add_heading(section_key, 0)
        doc.add_paragraph(response)
        path = f'Knowldgebase/SD/{section_key}{projectid}.docx'
        # Save the document
        doc.save(path)
        # Append the response to the project_info with a newline
        # if project_details:
        #    project_details += f"\n{section_key}:\n{response}"
        # else:
        #    project_details = f"{section_key}:\n{response}"
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document

        # doc = Document()
        # doc.add_heading('Static Website', 0)

        # for section in sections_queries.keys():
        # doc.add_heading(section.capitalize(), level=1)
        # doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        # return f"SRS document saved to {output_path}"
        query = "Generate the detailed static website styles.css code for each html provided in given document. this is the ui requirments.Ensure that the generated code strictly adheres to the information provided without adding any extra content or features not mentioned in the document. Provide detailed and clear answers, focusing solely on the specified requirements."
        path = output_path
        #        sections_queries = {
        #          "html": {
        #              "query": f"Write static html single page website with (.html files,style.css,script.js) code for the project based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
        #            "path": file_path}}
        project_info = {}
        #        project_details = ""
        #        for section, details in sections_queries.items():
        print("query", query)
        # query = f"For Project {projectname}, {details.get('query')}"
        response = self.generate_response_from_document(path, query, project_info)
        # if section == "Django Directory":
        #    self.create_directory_structure_from_text(response, base_dir)
        # Ensure the section key is formatted correctly
        section_key = "styles"
        # project_info[section_key] = response
        # Create a new .docx document
        doc = Document()
        doc.add_heading(section_key, 0)
        doc.add_paragraph(response)
        path = f'Knowldgebase/SD/{section_key}_{projectid}.docx'
        # Save the document
        doc.save(path)
        # Append the response to the project_info with a newline
        # if project_details:
        #    project_details += f"\n{section_key}:\n{response}"
        # else:
        #    project_details = f"{section_key}:\n{response}"
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document

        # doc = Document()
        # doc.add_heading('Static Website', 0)

        # for section in sections_queries.keys():
        # doc.add_heading(section.capitalize(), level=1)
        # doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(path)
        return "Code has been generated."


agent = SoftwareDeveloperAgent()
#projectname="Online Ecommerce Website on Django Platform"
output_path = 'htmloutput.docx'
base_dir= 'outputhtml'
os.makedirs(base_dir, exist_ok=True)
input_path='Knowldgebase/SA/SRS{projectid}.docx'
#response = agent.software_developer_agent_full(input_path, output_path,base_dir,projectid)
response = agent.software_developer_agent_template(input_path, output_path, base_dir,projectid,ui_data)